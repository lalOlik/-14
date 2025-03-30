#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Калькулятор Производных

Программа для определения производных сложных функций с возможностью 
сохранения результатов в различных форматах.
"""

import os
import sys
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import sympy as sp
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
from sympy.printing.mathml import mathml
from sympy.printing.latex import latex
import pandas as pd
from docx import Document
from docx.shared import Inches
import io
from scipy.io import savemat
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# Определяем символьные переменные для использования в выражениях
x, y, z = sp.symbols('x y z')
t = sp.symbols('t')
a, b, c, n, m = sp.symbols('a b c n m', constant=True)

class DerivativeCalculator:
    """Основной класс калькулятора производных"""
    
    def __init__(self):
        """Инициализация калькулятора"""
        # База данных табличных производных
        self.base_derivatives = {
            'sin(x)': sp.cos(x),
            'cos(x)': -sp.sin(x),
            'tan(x)': sp.sec(x)**2,
            'cot(x)': -sp.csc(x)**2,
            'sec(x)': sp.sec(x) * sp.tan(x),
            'csc(x)': -sp.csc(x) * sp.cot(x),
            'exp(x)': sp.exp(x),
            'log(x)': 1/x,
            'ln(x)': 1/x,
            'asin(x)': 1/sp.sqrt(1 - x**2),
            'acos(x)': -1/sp.sqrt(1 - x**2),
            'atan(x)': 1/(1 + x**2),
            'acot(x)': -1/(1 + x**2),
            'sinh(x)': sp.cosh(x),
            'cosh(x)': sp.sinh(x),
            'tanh(x)': 1 - sp.tanh(x)**2,
            'coth(x)': 1 - sp.coth(x)**2,
            'x^n': n*x**(n-1),
            'a^x': sp.log(a) * a**x,
            'sqrt(x)': 1/(2*sp.sqrt(x)),
        }
        # История вычислений
        self.history = []
    
    def get_derivative(self, expr_str, var='x', order=1):
        """Вычисление производной выражения"""
        try:
            # Преобразуем строку в sympy выражение
            expr = sp.sympify(expr_str)
            
            # Определяем переменную дифференцирования
            diff_var = sp.Symbol(var)
            
            # Вычисляем производную нужного порядка
            result = sp.diff(expr, diff_var, order)
            
            # Упрощаем результат
            simplified = self.simplify_expression(result)
            
            # Добавляем в историю
            self.history.append({
                'expression': expr_str,
                'variable': var,
                'order': order,
                'result': simplified
            })
            
            return simplified
        except Exception as e:
            return f"Ошибка: {str(e)}"
    
    def get_partial_derivative(self, expr_str, variables, orders):
        """Вычисление частной производной выражения"""
        try:
            # Преобразуем строку в sympy выражение
            expr = sp.sympify(expr_str)
            
            # Вычисляем частную производную
            result = expr
            for var, order in zip(variables, orders):
                diff_var = sp.Symbol(var)
                result = sp.diff(result, diff_var, order)
            
            # Упрощаем результат
            simplified = self.simplify_expression(result)
            
            # Добавляем в историю
            self.history.append({
                'expression': expr_str,
                'variables': variables,
                'orders': orders,
                'result': simplified
            })
            
            return simplified
        except Exception as e:
            return f"Ошибка: {str(e)}"
    
    def simplify_expression(self, expr):
        """Упрощение выражения"""
        try:
            # Применяем различные методы упрощения
            simplified = sp.simplify(expr)
            simplified = sp.trigsimp(simplified)
            simplified = sp.expand(simplified)
            simplified = sp.factor(simplified)
            simplified = sp.collect(simplified, x)
            
            return simplified
        except:
            return expr
    
    def save_result(self, result, filename, format_type):
        """Сохранение результата в файл"""
        try:
            if format_type == 'pdf':
                self._save_as_pdf(result, filename)
            elif format_type == 'doc':
                self._save_as_doc(result, filename)
            elif format_type == 'rtf':
                self._save_as_rtf(result, filename)
            elif format_type == 'mat':
                self._save_as_mat(result, filename)
            return True
        except Exception as e:
            print(f"Ошибка при сохранении: {str(e)}")
            return False
    
    def _save_as_pdf(self, result, filename):
        """Сохранение в PDF формат"""
        c = canvas.Canvas(filename, pagesize=letter)
        c.setFont("Helvetica", 12)
        c.drawString(100, 750, "Результат дифференцирования:")
        c.drawString(100, 730, str(result))
        
        # Если есть история, добавляем ее
        if self.history:
            c.drawString(100, 700, "История вычислений:")
            y_pos = 680
            for i, entry in enumerate(self.history[-5:]):  # Последние 5 записей
                c.drawString(120, y_pos, f"{i+1}. {entry['expression']} -> {entry['result']}")
                y_pos -= 20
        c.save()
    
    def _save_as_doc(self, result, filename):
        """Сохранение в DOC формат"""
        doc = Document()
        doc.add_heading('Результат дифференцирования', 0)
        doc.add_paragraph(str(result))
        
        # Если есть история, добавляем ее
        if self.history:
            doc.add_heading('История вычислений', level=1)
            for i, entry in enumerate(self.history):
                doc.add_paragraph(f"{i+1}. {entry['expression']} -> {entry['result']}")
        
        doc.save(filename)
    
    def _save_as_rtf(self, result, filename):
        """Сохранение в RTF формат"""
        # Используем docx и сохраняем в .rtf
        doc = Document()
        doc.add_heading('Результат дифференцирования', 0)
        doc.add_paragraph(str(result))
        
        # Если есть история, добавляем ее
        if self.history:
            doc.add_heading('История вычислений', level=1)
            for i, entry in enumerate(self.history):
                doc.add_paragraph(f"{i+1}. {entry['expression']} -> {entry['result']}")
        
        doc.save(filename)
    
    def _save_as_mat(self, result, filename):
        """Сохранение в MATLAB формат"""
        # Преобразуем результат в строку для сохранения
        result_str = str(result)
        savemat(filename, {'result': result_str, 'history': str(self.history)})


class DerivativeCalculatorUI:
    """Пользовательский интерфейс для калькулятора производных"""
    
    def __init__(self):
        """Инициализация пользовательского интерфейса"""
        self.calculator = DerivativeCalculator()
        self.root = tk.Tk()
        self.root.title("Калькулятор Производных")
        self.root.geometry("800x600")
        
        # Настройка темы
        style = ttk.Style()
        style.theme_use('clam')
        
        self._create_widgets()
    
    def _create_widgets(self):
        """Создание элементов интерфейса"""
        # Создаем фреймы
        input_frame = ttk.LabelFrame(self.root, text="Ввод выражения")
        input_frame.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)
        
        result_frame = ttk.LabelFrame(self.root, text="Результат")
        result_frame.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)
        
        # Элементы ввода
        ttk.Label(input_frame, text="Выражение:").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        self.expr_entry = ttk.Entry(input_frame, width=40)
        self.expr_entry.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W+tk.E)
        
        # Примеры выражений
        ttk.Label(input_frame, text="Примеры: sin(x), x^2 + 3*x, exp(x)*cos(x)").grid(row=1, column=1, padx=5, pady=5, sticky=tk.W)
        
        # Выбор переменной
        ttk.Label(input_frame, text="Переменная:").grid(row=2, column=0, padx=5, pady=5, sticky=tk.W)
        self.var_combobox = ttk.Combobox(input_frame, values=['x', 'y', 'z', 't'])
        self.var_combobox.grid(row=2, column=1, padx=5, pady=5, sticky=tk.W)
        self.var_combobox.current(0)
        
        # Порядок производной
        ttk.Label(input_frame, text="Порядок:").grid(row=3, column=0, padx=5, pady=5, sticky=tk.W)
        self.order_spinbox = ttk.Spinbox(input_frame, from_=1, to=10, width=5)
        self.order_spinbox.grid(row=3, column=1, padx=5, pady=5, sticky=tk.W)
        
        # Кнопки
        buttons_frame = ttk.Frame(input_frame)
        buttons_frame.grid(row=4, column=0, columnspan=2, padx=5, pady=10)
        
        ttk.Button(buttons_frame, text="Вычислить", command=self._calculate_derivative).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="Частная производная", command=self._calculate_partial_derivative).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="Упростить", command=self._simplify).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="Сохранить", command=self._save_result).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="Очистить", command=self._clear).pack(side=tk.LEFT, padx=5)
        
        # Результаты
        self.result_text = tk.Text(result_frame, height=15, wrap=tk.WORD)
        self.result_text.pack(padx=5, pady=5, fill=tk.BOTH, expand=True)
        
        # Добавляем скроллбар для текстового поля
        scrollbar = ttk.Scrollbar(self.result_text, command=self.result_text.yview)
        self.result_text.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Табличные производные
        self._add_table_tab()
    
    def _add_table_tab(self):
        """Добавление вкладки с табличными производными"""
        # Создаем фрейм для таблицы
        table_frame = ttk.LabelFrame(self.root, text="Табличные производные")
        table_frame.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)
        
        # Создаем таблицу
        table = ttk.Treeview(table_frame, columns=("Function", "Derivative"), show="headings")
        table.heading("Function", text="Функция")
        table.heading("Derivative", text="Производная")
        table.column("Function", width=100)
        table.column("Derivative", width=200)
        
        # Заполняем таблицу данными
        for func, deriv in self.calculator.base_derivatives.items():
            table.insert("", tk.END, values=(func, str(deriv)))
        
        # Добавляем скроллбар для таблицы
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, command=table.yview)
        table.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        table.pack(padx=5, pady=5, fill=tk.BOTH, expand=True)
    
    def _calculate_derivative(self):
        """Вычисление производной"""
        expr = self.expr_entry.get()
        var = self.var_combobox.get()
        order = int(self.order_spinbox.get())
        
        if not expr:
            messagebox.showerror("Ошибка", "Введите выражение")
            return
        
        result = self.calculator.get_derivative(expr, var, order)
        self._display_result(f"Производная {order}-го порядка от {expr} по {var}:\n\n{result}")
    
    def _calculate_partial_derivative(self):
        """Вычисление частной производной"""
        expr = self.expr_entry.get()
        
        if not expr:
            messagebox.showerror("Ошибка", "Введите выражение")
            return
        
        # Открываем диалоговое окно для ввода переменных и порядков
        partial_window = tk.Toplevel(self.root)
        partial_window.title("Частная производная")
        partial_window.geometry("400x300")
        
        vars_frame = ttk.Frame(partial_window)
        vars_frame.pack(padx=10, pady=10, fill=tk.BOTH)
        
        ttk.Label(vars_frame, text="Переменные (через запятую):").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        vars_entry = ttk.Entry(vars_frame, width=20)
        vars_entry.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W+tk.E)
        vars_entry.insert(0, "x, y")
        
        ttk.Label(vars_frame, text="Порядки (через запятую):").grid(row=1, column=0, padx=5, pady=5, sticky=tk.W)
        orders_entry = ttk.Entry(vars_frame, width=20)
        orders_entry.grid(row=1, column=1, padx=5, pady=5, sticky=tk.W+tk.E)
        orders_entry.insert(0, "1, 1")
        
        def on_calculate():
            vars_str = vars_entry.get()
            orders_str = orders_entry.get()
            
            variables = [v.strip() for v in vars_str.split(",")]
            orders = [int(o.strip()) for o in orders_str.split(",")]
            
            if len(variables) != len(orders):
                messagebox.showerror("Ошибка", "Количество переменных и порядков должно совпадать")
                return
            
            result = self.calculator.get_partial_derivative(expr, variables, orders)
            
            var_orders = ", ".join([f"∂{o}{v}" for v, o in zip(variables, orders)])
            self._display_result(f"Частная производная {var_orders} от {expr}:\n\n{result}")
            partial_window.destroy()
        
        ttk.Button(vars_frame, text="Вычислить", command=on_calculate).grid(row=2, column=0, columnspan=2, padx=5, pady=10)
    
    def _simplify(self):
        """Упрощение выражения"""
        expr = self.expr_entry.get()
        
        if not expr:
            messagebox.showerror("Ошибка", "Введите выражение")
            return
        
        try:
            # Преобразуем строку в sympy выражение
            sympy_expr = sp.sympify(expr)
            
            # Упрощаем выражение
            result = self.calculator.simplify_expression(sympy_expr)
            
            self._display_result(f"Упрощенное выражение для {expr}:\n\n{result}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось упростить выражение: {str(e)}")
    
    def _save_result(self):
        """Сохранение результата в файл"""
        # Получаем текущий результат
        result_text = self.result_text.get("1.0", tk.END).strip()
        
        if not result_text:
            messagebox.showerror("Ошибка", "Нет результата для сохранения")
            return
        
        # Открываем диалоговое окно выбора формата и места сохранения
        file_types = [
            ("PDF файлы", "*.pdf"),
            ("DOC файлы", "*.docx"),
            ("RTF файлы", "*.rtf"),
            ("MATLAB файлы", "*.mat"),
        ]
        filename = filedialog.asksaveasfilename(
            filetypes=file_types,
            defaultextension=".pdf"
        )
        
        if filename:
            # Определяем формат по расширению
            extension = os.path.splitext(filename)[1][1:].lower()
            format_map = {
                "pdf": "pdf",
                "docx": "doc",
                "rtf": "rtf",
                "mat": "mat"
            }
            
            format_type = format_map.get(extension, "pdf")
            
            # Сохраняем результат
            try:
                success = self.calculator.save_result(result_text, filename, format_type)
                if success:
                    messagebox.showinfo("Успех", f"Результат сохранен в {filename}")
                else:
                    messagebox.showerror("Ошибка", "Не удалось сохранить результат")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка при сохранении: {str(e)}")
    
    def _clear(self):
        """Очистка полей ввода и результата"""
        self.expr_entry.delete(0, tk.END)
        self.var_combobox.current(0)
        self.order_spinbox.delete(0, tk.END)
        self.order_spinbox.insert(0, "1")
        self.result_text.delete("1.0", tk.END)
    
    def _display_result(self, text):
        """Отображение результата в текстовом поле"""
        self.result_text.delete("1.0", tk.END)
        self.result_text.insert(tk.END, text)
    
    def run(self):
        """Запуск приложения"""
        self.root.mainloop()


if __name__ == "__main__":
    app = DerivativeCalculatorUI()
    app.run() 
